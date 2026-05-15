"""
博客系统 - Flask 后端
使用 SQLite 数据库，支持文章管理
"""

from flask import Flask, request, jsonify, render_template, redirect, url_for, session
from flask import send_from_directory
import sqlite3
from datetime import datetime
import os
import hashlib
import markdown
import uuid
from io import BytesIO

# 新增：Word文档处理
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re

# 版本号 - 更新此值可强制刷新浏览器缓存
VERSION = '1.1.3'

app = Flask(__name__)
app.secret_key = 'blog-secret-key-2026'
app.config['DATABASE'] = 'blog.db'
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50MB

# 确保上传目录存在
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# ============================================
# 数据库工具函数
# ============================================

def get_db():
    """获取数据库连接"""
    conn = sqlite3.connect(app.config['DATABASE'])
    conn.row_factory = sqlite3.Row
    return conn

def init_db():
    """初始化数据库"""
    conn = get_db()
    cursor = conn.cursor()
    
    # 创建文章表
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS posts (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            title TEXT NOT NULL,
            slug TEXT UNIQUE NOT NULL,
            excerpt TEXT,
            content TEXT NOT NULL,
            tags TEXT,
            author TEXT DEFAULT '觉醒者',
            views INTEGER DEFAULT 0,
            featured INTEGER DEFAULT 0,
            status TEXT DEFAULT 'published',
            created_at TEXT DEFAULT CURRENT_TIMESTAMP,
            updated_at TEXT DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    
    # 创建标签表
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS tags (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT UNIQUE NOT NULL,
            slug TEXT UNIQUE NOT NULL
        )
    ''')
    
    # 创建文章-标签关联表
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS post_tags (
            post_id INTEGER,
            tag_id INTEGER,
            PRIMARY KEY (post_id, tag_id),
            FOREIGN KEY (post_id) REFERENCES posts(id),
            FOREIGN KEY (tag_id) REFERENCES tags(id)
        )
    ''')
    
    # 创建管理员表
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS admin (
            id INTEGER PRIMARY KEY,
            username TEXT UNIQUE NOT NULL,
            password TEXT NOT NULL
        )
    ''')
    
    # 检查是否需要初始化默认管理员
    cursor.execute('SELECT * FROM admin WHERE username = ?', ('admin',))
    if not cursor.fetchone():
        # 默认管理员密码: admin123 (SHA256哈希)
        default_password = hashlib.sha256('admin123'.encode()).hexdigest()
        cursor.execute('INSERT INTO admin (username, password) VALUES (?, ?)', 
                      ('admin', default_password))
    
    # 创建关于页面表
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS about_page (
            id INTEGER PRIMARY KEY,
            content TEXT NOT NULL,
            updated_at TEXT DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    
    # 创建站点配置表
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS site_settings (
            key TEXT PRIMARY KEY,
            value TEXT NOT NULL
        )
    ''')
    
    # 初始化默认站点配置
    default_settings = {
        'site_name': '认知觉醒',
        'site_title': '认知觉醒',
        'site_subtitle': '探索思维边界，觉醒生命潜能',
        'site_description': '专注于个人成长、心理学与认知科学的博客',
        'site_keywords': '认知觉醒,个人成长,心理学,思维,觉醒',
        'site_logo': '☯',
        'hero_title': '认知觉醒',
        'about_title': '关于「认知觉醒」',
        'about_subtitle': '探索思维边界，觉醒生命潜能',
        'about_avatar': '☯',
        'footer_slogan': '探索思维边界，觉醒生命潜能',
        'footer_copyright': '© 2026 认知觉醒 · 用心原创，欢迎分享',
        'contact_email': 'awakening@example.com',
        'contact_wechat': '认知觉醒（待开通）'
    }
    
    for key, value in default_settings.items():
        cursor.execute('INSERT OR IGNORE INTO site_settings (key, value) VALUES (?, ?)', (key, value))
    
    # 初始化默认关于页面内容
    cursor.execute('SELECT * FROM about_page WHERE id = 1')
    if not cursor.fetchone():
        default_about = '''## 🌱 创办初衷

在信息爆炸的时代，我们每天被动接收大量碎片化的信息，却很少停下来思考：

- 我们为什么会这样想？
- 我们为什么会有这样的情绪？
- 我们如何才能更好地认识自己？

创办这个博客，是希望能够：

1. **记录思考** — 将阅读和学习中的思考整理成文
2. **分享知识** — 把有用的认知科学知识用通俗易懂的方式分享
3. **共同成长** — 与志同道合的朋友一起探索成长的奥秘

## 📚 内容方向

- **认知科学** — 了解大脑如何运作，思维是如何形成的
- **心理学** — 理解人类行为和心理机制的深层原因
- **个人成长** — 探索自我提升的方法与实践
- **实践心得** — 将理论付诸实践的经验总结

> 觉醒不是终点，而是一场永无止境的旅程。

## 👤 关于作者

一个持续探索自我成长之路的普通人。

相信每个人都具备觉醒的潜能，只是在等待一个契机。

希望通过这个博客，与更多志同道合的朋友相遇，共同在认知的道路上不断前进。

## 📬 联系方式

如果你有任何想法或建议，欢迎通过以下方式联系我：

- 📧 邮箱：awakening@example.com
- 💬 微信公众号：认知觉醒（待开通）

---

*感谢你的阅读，愿我们都能在认知的道路上不断觉醒。*
'''
        cursor.execute('INSERT INTO about_page (id, content) VALUES (1, ?)', (default_about,))
    
    # 插入示例文章
    cursor.execute('SELECT COUNT(*) FROM posts')
    if cursor.fetchone()[0] == 0:
        insert_sample_posts(conn)
    
    conn.commit()
    conn.close()

def insert_sample_posts(conn):
    """插入示例文章"""
    cursor = conn.cursor()
    
    sample_posts = [
        {
            'title': '认知觉醒：从混沌到清醒的旅程',
            'slug': 'awakening-journey',
            'excerpt': '我们终其一生都在追求觉醒。觉醒不是突然的顿悟，而是一个渐进的过程——从无意识的自动化反应，到有意识的觉察与选择。',
            'content': '''# 认知觉醒：从混沌到清醒的旅程

我们终其一生都在追求觉醒。但什么才是真正的觉醒？

## 觉醒的三个层次

### 第一层：察觉
- 意识到自己的思维模式
- 看见自己的情绪波动
- 觉察自己的行为习惯

> "当你能够观察自己的思维时，你就已经开始觉醒了。"

### 第二层：理解
- 理解思维背后的运作机制
- 明白情绪产生的根源
- 认识习惯形成的路径

### 第三层：转化
- 主动改变思维模式
- 转化负面情绪为动力
- 重塑新的行为习惯

## 如何开始觉醒

1. **每日反思** - 留出时间回顾一天的行为和思维
2. **冥想练习** - 培养专注力和觉察力
3. **阅读经典** - 站在巨人的肩膀上理解世界
4. **践行验证** - 用实践检验认知的正确性

觉醒不是终点，而是一场永无止境的旅程。每一个当下都是觉醒的机会。''',
            'tags': '觉醒,成长,心理学',
            'featured': 1
        },
        {
            'title': '大脑的默认模式网络：为什么我们总是走神',
            'slug': 'default-mode-network',
            'excerpt': '当你坐在桌前准备工作时，大脑却在自动播放"我晚饭吃什么"——这不是你的意志力问题，而是大脑默认模式网络在作祟。',
            'content': '''# 大脑的默认模式网络：为什么我们总是走神

你是否有过这样的体验：

- 坐在电脑前准备写报告，思绪却飘到了晚餐
- 开会时身体在场，灵魂却在神游
- 明明在听别人说话，脑子里却在想别的事

这不是你不够专注，而是你的大脑在执行它的"默认程序"。

## 什么是默认模式网络（DMN）

当大脑不专注于外部任务时，就会激活一组特定的脑区，这些脑区形成的网络被称为**默认模式网络（Default Mode Network, DMN）**。

### DMN 的主要功能

1. **内心独白** - 自言自语、内心对话
2. **过去未来** - 回忆过去、畅想未来
3. **他人视角** - Theory of Mind，理解他人想法
4. **创意发散** - 产生创意和灵感

## DMN 的双面性

### 积极的 DMN
- 促进创意和创造力
- 帮助整合记忆
- 支持社会认知

### 消极的 DMN
- 导致拖延和走神
- 产生焦虑和担忧
- 陷入负面思维循环

## 驯服 DMN 的方法

| 方法 | 作用 |
|------|------|
| 正念冥想 | 增强对思维的觉察 |
| 设定明确目标 | 给大脑明确的执行任务 |
| 番茄工作法 | 主动管理注意力 |
| 身体扫描 | 将注意力带回当下 |

> 关键不是阻止 DMN，而是学会在需要时切换。

理解大脑的运作机制，不是为了对抗它，而是为了更好地与它合作。''',
            'tags': '大脑,专注力,认知科学',
            'featured': 1
        },
        {
            'title': '心流体验：进入最佳状态的秘诀',
            'slug': 'flow-state',
            'excerpt': '那种全神贯注、时间消失、完全沉浸在活动中的感觉，就是心流。学会触发心流，就等于掌握了进入最佳状态的钥匙。',
            'content': '''# 心流体验：进入最佳状态的秘诀

你有没有过这种体验：

- 打球时感觉不到时间流逝
- 画画时忘记了周围的一切
- 编程时进入"人机合一"的境界

这就是心理学家米哈里·契克森米哈伊提出的**心流（Flow）**状态。

## 心流的特征

1. **全神贯注** - 注意力完全集中在当下
2. **行动与意识合一** - 思想和行动无缝连接
3. **失去自我意识** - 不再担心别人的看法
4. **时间扭曲** - 时间感觉变快或变慢
5. **即时反馈** - 清晰的行动结果反馈
6. **控制感** - 感觉一切尽在掌控

## 进入心流的条件

### 挑战与技能的平衡

```
难度太高 → 焦虑
难度太低 → 无聊
难度适中 → 心流
```

### 心流触发器

- **清晰的目标** - 知道要做什么
- **即时反馈** - 知道做得好不好
- **全神贯注** - 排除干扰
- **控制感** - 感觉能够掌控

## 如何培养心流能力

1. **选择热爱的事** - 心流更容易在感兴趣的事上出现
2. **设定具体目标** - 明确要达成什么
3. **消除干扰** - 手机静音，关闭通知
4. **从小开始** - 先从短时间的心流练习

> 心流不是天才的专利，而是每个人都可以学习的技能。

当你找到那个挑战与能力完美匹配的甜蜜点，心流就会自然涌现。''',
            'tags': '心流,专注力,效率',
            'featured': 0
        },
        {
            'title': '情绪的隐喻：为什么我们需要学会命名',
            'slug': 'emotion-naming',
            'excerpt': '"你今天感觉怎么样？"——我们常常不知道如何回答。情绪粒度低的人，往往被情绪所控制；而情绪粒度高的人，却能巧妙地与情绪共处。',
            'content': '''# 情绪的隐喻：为什么我们需要学会命名

你有没有过这种感觉：

- 心里很不舒服，但说不清是什么
- 明明是生气，却表现得很冷漠
- 感到"低落"，却不知道是悲伤还是疲惫

这可能是你的**情绪粒度（Emotional Granularity）**较低。

## 什么是情绪粒度

情绪粒度是指个体区分和命名不同情绪感受的能力。

### 高情绪粒度 vs 低情绪粒度

| 高情绪粒度 | 低情绪粒度 |
|-----------|-----------|
| 能精确区分：失望、沮丧、绝望 | 统一称为"难受" |
| 能识别：愤怒、恼怒、烦躁 | 统一称为"不爽" |
| 能分辨：感动、温暖、感激 | 统一称为"心里暖暖的" |

## 为什么命名情绪如此重要

### 1. 命名即疗愈

> 当我们能够准确命名一种情绪时，它对我们的影响就会减弱。

这被称为**情绪标注效应（Effect of Emotion Labeling）**。

### 2. 增强情绪调节能力

- 能够区分情绪，才能针对性地应对
- 把"我很焦虑"变成"我对这件事有些担忧，但我在掌控中"

### 3. 改善人际关系

- 准确表达情绪，让他人更好地理解你
- "我感到有些失落"比"你让我很失望"更有效

## 提升情绪粒度的方法

1. **情绪词汇表** - 学习更多情绪词汇
2. **情绪日记** - 每天记录情绪和触发因素
3. **反思练习** - 问自己"我现在具体是什么感受？"
4. **阅读文学** - 小说是很好的情绪教科书

情绪不是敌人，而是信使。学会解读它们的信息，我们就能更好地理解自己。''',
            'tags': '情绪,情商,心理学',
            'featured': 0
        },
        {
            'title': '自我决定理论：什么才是真正的内在动机',
            'slug': 'self-determination-theory',
            'excerpt': '外部奖励有时会削弱内在动机。理解自我决定理论的三个基本心理需求，让你重新找回做事的动力。',
            'content': '''# 自我决定理论：什么才是真正的内在动机

为什么有些人工作像打了鸡血，有些人却总是提不起劲？

心理学家爱德华·德西和理查德·瑞安提出的**自我决定理论（Self-Determination Theory, SDT）**给了我们答案。

## 人类的三种基本心理需求

### 1. 自主感（Autonomy）

> 我选择这样做，而不是被迫这样做。

- 能够自己决定做什么、怎么做
- 感受到行为的自我一致性
- 拥有选择权

### 2. 胜任感（Competence）

> 我能够做到，并且做得越来越好。

- 感受到挑战与技能的匹配
- 获得成长和进步的体验
- 得到有意义的反馈

### 3. 归属感（Relatedness）

> 我被看见，被理解，与他人连接。

- 感受到与他人的情感联结
- 体验到被支持、被接纳
- 拥有归属的社群

## 内在动机 vs 外在动机

### 外在动机的层次

1. **外部调节** - 完全受外部奖惩控制
2. **内射调节** - 为了避免愧疚/获得自尊
3. **认同调节** - 认同目标的价值
4. **整合调节** - 与自我完全整合

### 内在动机

当三种基本需求都得到满足时，外在动机可以转化为内在动机。

## 如何激发内在动机

1. **提供选择** - 即使是小选择也能增加自主感
2. **设定适度挑战** - 难度太高会挫折，太低会无聊
3. **给予自主支持** - 解释"为什么"而非只说"做什么"
4. **创造连接** - 建立支持性的人际关系

> 真正的动力来自内心，而非外部奖励。

当你做的事情满足了你的自主、胜任和归属需求时，你会发现自己根本不需要"坚持"——因为这已经成为一种享受。''',
            'tags': '动机,心理学,成长',
            'featured': 1
        }
    ]
    
    for post in sample_posts:
        cursor.execute('''
            INSERT INTO posts (title, slug, excerpt, content, tags, featured, created_at, updated_at)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?)
        ''', (post['title'], post['slug'], post['excerpt'], post['content'], 
              post['tags'], post['featured'], datetime.now().isoformat(), datetime.now().isoformat()))
        
        post_id = cursor.lastrowid
        
        # 处理标签
        for tag_name in post['tags'].split(','):
            tag_name = tag_name.strip()
            tag_slug = tag_name.lower()
            
            # 插入或获取标签
            cursor.execute('INSERT OR IGNORE INTO tags (name, slug) VALUES (?, ?)', (tag_name, tag_slug))
            cursor.execute('SELECT id FROM tags WHERE slug = ?', (tag_slug,))
            tag_row = cursor.fetchone()
            if tag_row:
                cursor.execute('INSERT OR IGNORE INTO post_tags (post_id, tag_id) VALUES (?, ?)', 
                             (post_id, tag_row['id']))

# ============================================
# 认证相关
# ============================================

def verify_admin(username, password):
    """验证管理员"""
    conn = get_db()
    cursor = conn.cursor()
    password_hash = hashlib.sha256(password.encode()).hexdigest()
    cursor.execute('SELECT * FROM admin WHERE username = ? AND password = ?', 
                  (username, password_hash))
    result = cursor.fetchone()
    conn.close()
    return result is not None

def login_required(f):
    """登录装饰器"""
    from functools import wraps
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'admin' not in session:
            return redirect(url_for('admin_login'))
        return f(*args, **kwargs)
    return decorated_function

# ============================================
# API 路由 - 公开接口
# ============================================

@app.route('/api/posts', methods=['GET'])
def get_posts():
    """获取文章列表"""
    conn = get_db()
    cursor = conn.cursor()
    
    tag = request.args.get('tag')
    status = request.args.get('status', 'published')
    
    if tag:
        cursor.execute('''
            SELECT p.* FROM posts p
            JOIN post_tags pt ON p.id = pt.post_id
            JOIN tags t ON pt.tag_id = t.id
            WHERE t.name = ? AND p.status = ?
            ORDER BY p.created_at DESC
        ''', (tag, status))
    else:
        cursor.execute('''
            SELECT * FROM posts WHERE status = ?
            ORDER BY featured DESC, created_at DESC
        ''', (status,))
    
    posts = [dict(row) for row in cursor.fetchall()]
    conn.close()
    
    # 转换标签为列表
    for post in posts:
        if post.get('tags'):
            post['tags'] = [t.strip() for t in post['tags'].split(',')]
    
    return jsonify(posts)

@app.route('/api/posts/<slug>', methods=['GET'])
def get_post(slug):
    """获取单篇文章"""
    conn = get_db()
    cursor = conn.cursor()
    cursor.execute('SELECT * FROM posts WHERE slug = ?', (slug,))
    post = cursor.fetchone()
    
    if post:
        # 增加阅读量
        cursor.execute('UPDATE posts SET views = views + 1 WHERE slug = ?', (slug,))
        conn.commit()
        post = dict(post)
        if post.get('tags'):
            post['tags'] = [t.strip() for t in post['tags'].split(',')]
    else:
        post = None
    
    conn.close()
    return jsonify(post) if post else (jsonify({'error': '文章不存在'}), 404)

@app.route('/api/tags', methods=['GET'])
def get_tags():
    """获取所有标签"""
    conn = get_db()
    cursor = conn.cursor()
    cursor.execute('''
        SELECT t.*, COUNT(pt.post_id) as count 
        FROM tags t
        LEFT JOIN post_tags pt ON t.id = pt.tag_id
        LEFT JOIN posts p ON pt.post_id = p.id AND p.status = 'published'
        GROUP BY t.id
        ORDER BY count DESC
    ''')
    tags = [dict(row) for row in cursor.fetchall()]
    conn.close()
    return jsonify(tags)

# ============================================
# API 路由 - 站点配置接口
# ============================================

def get_settings():
    """获取站点配置字典"""
    conn = get_db()
    cursor = conn.cursor()
    cursor.execute('SELECT key, value FROM site_settings')
    settings = {row['key']: row['value'] for row in cursor.fetchall()}
    conn.close()
    return settings

@app.route('/api/settings', methods=['GET'])
def get_settings_api():
    """获取站点配置（API 接口）"""
    return jsonify(get_settings())

# ============================================
# API 路由 - 关于页面接口
# ============================================

@app.route('/api/about', methods=['GET'])
def get_about():
    """获取关于页面内容"""
    conn = get_db()
    cursor = conn.cursor()
    cursor.execute('SELECT content, updated_at FROM about_page WHERE id = 1')
    row = cursor.fetchone()
    conn.close()
    
    if row:
        return jsonify({
            'content': row['content'],
            'updated_at': row['updated_at']
        })
    return jsonify({'content': '', 'updated_at': None})

# ============================================
# API 路由 - 管理接口
# ============================================

@app.route('/api/admin/about', methods=['PUT'])
@login_required
def update_about():
    """更新关于页面内容"""
    data = request.get_json()
    content = data.get('content', '')
    
    conn = get_db()
    cursor = conn.cursor()
    now = datetime.now().isoformat()
    
    cursor.execute('''
        UPDATE about_page SET content = ?, updated_at = ? WHERE id = 1
    ''', (content, now))
    
    conn.commit()
    conn.close()
    
    return jsonify({'success': True})

@app.route('/api/admin/settings', methods=['PUT'])
@login_required
def update_settings():
    """更新站点配置"""
    data = request.get_json()
    
    conn = get_db()
    cursor = conn.cursor()
    
    for key, value in data.items():
        cursor.execute('''
            INSERT OR REPLACE INTO site_settings (key, value) VALUES (?, ?)
        ''', (key, value))
    
    conn.commit()
    conn.close()
    
    return jsonify({'success': True})

@app.route('/api/admin/login', methods=['POST'])
def admin_login_api():
    """管理员登录"""
    data = request.get_json()
    username = data.get('username')
    password = data.get('password')
    
    if verify_admin(username, password):
        session['admin'] = username
        return jsonify({'success': True, 'message': '登录成功'})
    return jsonify({'success': False, 'message': '用户名或密码错误'}), 401

@app.route('/api/admin/logout', methods=['POST'])
def admin_logout_api():
    """管理员登出"""
    session.pop('admin', None)
    return jsonify({'success': True})

@app.route('/api/admin/password', methods=['PUT'])
@login_required
def update_password():
    """修改密码"""
    data = request.get_json()
    old_password = data.get('oldPassword')
    new_password = data.get('newPassword')
    
    if not old_password or not new_password:
        return jsonify({'success': False, 'message': '请填写完整信息'}), 400
    
    if len(new_password) < 6:
        return jsonify({'success': False, 'message': '新密码长度不能少于6位'}), 400
    
    conn = get_db()
    cursor = conn.cursor()
    
    # 验证旧密码
    old_hash = hashlib.sha256(old_password.encode()).hexdigest()
    cursor.execute('SELECT * FROM admin WHERE password = ?', (old_hash,))
    if not cursor.fetchone():
        conn.close()
        return jsonify({'success': False, 'message': '原密码错误'}), 401
    
    # 更新密码
    new_hash = hashlib.sha256(new_password.encode()).hexdigest()
    cursor.execute('UPDATE admin SET password = ?', (new_hash,))
    conn.commit()
    conn.close()
    
    return jsonify({'success': True, 'message': '密码修改成功'})

@app.route('/api/admin/username', methods=['PUT'])
@login_required
def update_username():
    """修改用户名"""
    data = request.get_json()
    new_username = data.get('username')
    password = data.get('password')
    
    if not new_username or not password:
        return jsonify({'success': False, 'message': '请填写完整信息'}), 400
    
    if len(new_username) < 3:
        return jsonify({'success': False, 'message': '用户名长度不能少于3位'}), 400
    
    conn = get_db()
    cursor = conn.cursor()
    
    # 验证密码
    password_hash = hashlib.sha256(password.encode()).hexdigest()
    cursor.execute('SELECT * FROM admin WHERE password = ?', (password_hash,))
    if not cursor.fetchone():
        conn.close()
        return jsonify({'success': False, 'message': '密码错误'}), 401
    
    # 检查新用户名是否已存在
    cursor.execute('SELECT * FROM admin WHERE username = ?', (new_username,))
    if cursor.fetchone():
        conn.close()
        return jsonify({'success': False, 'message': '用户名已存在'}), 400
    
    # 更新用户名
    cursor.execute('UPDATE admin SET username = ?', (new_username,))
    conn.commit()
    conn.close()
    
    # 更新 session
    session['admin'] = new_username
    
    return jsonify({'success': True, 'message': '用户名修改成功'})

@app.route('/api/admin/upload', methods=['POST'])
@login_required
def upload_image():
    """上传图片"""
    if 'file' not in request.files:
        return jsonify({'success': False, 'message': '未选择文件'}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({'success': False, 'message': '未选择文件'}), 400

    # 检查文件类型
    allowed = {'png', 'jpg', 'jpeg', 'gif', 'webp', 'svg'}
    ext = file.filename.rsplit('.', 1)[-1].lower() if '.' in file.filename else ''
    if ext not in allowed:
        return jsonify({'success': False, 'message': '不支持的图片格式'}), 400

    # 生成唯一文件名
    filename = f"{uuid.uuid4().hex}.{ext}"
    save_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(save_path)

    url = f"/uploads/{filename}"
    return jsonify({'success': True, 'url': url})


@app.route('/uploads/<path:filename>')
def serve_uploads(filename):
    """提供上传文件访问"""
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename)


@app.route('/api/admin/import-word', methods=['POST'])
@login_required
def import_word():
    """从Word文档导入文章内容"""
    if 'file' not in request.files:
        return jsonify({'success': False, 'message': '未选择文件'}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({'success': False, 'message': '未选择文件'}), 400

    # 检查文件类型
    ext = file.filename.rsplit('.', 1)[-1].lower() if '.' in file.filename else ''
    if ext not in {'docx'}:
        return jsonify({'success': False, 'message': '只支持 .docx 格式'}), 400

    try:
        # 获取文件名作为备选标题
        original_filename = file.filename.rsplit('.', 1)[0] if '.' in file.filename else ''
        file_title = re.sub(r'[-_]+', ' ', original_filename).strip()

        # 读取文件内容
        file_content = BytesIO(file.read())
        doc = Document(file_content)

        # 第一步：提取所有内联图片并保存
        saved_images = []
        for shape in doc.inline_shapes:
            try:
                # 获取图片的rId
                blip = shape._inline.graphic.graphicData.pic.blipFill.blip
                rId = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                # 通过rId获取图片数据
                image_part = doc.part.related_parts.get(rId)
                if image_part:
                    img_bytes = image_part.blob
                    # 获取扩展名
                    img_ext = 'png'
                    content_type = image_part.content_type if hasattr(image_part, 'content_type') else ''
                    if 'jpeg' in content_type or 'jpg' in content_type:
                        img_ext = 'jpg'
                    elif 'gif' in content_type:
                        img_ext = 'gif'
                    elif 'webp' in content_type:
                        img_ext = 'webp'

                    filename = f"{uuid.uuid4().hex}.{img_ext}"
                    image_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                    with open(image_path, 'wb') as f:
                        f.write(img_bytes)
                    saved_images.append(f"/uploads/{filename}")
                else:
                    saved_images.append(None)
            except Exception:
                saved_images.append(None)

        markdown_content = []
        title = file_title
        img_idx = 0

        # 第二步：处理所有body元素
        body_xml = doc.element.body

        for child in body_xml:
            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag

            if tag == 'p':
                # 段落
                para = next((p for p in doc.paragraphs if p._element is child), None)
                if para is None:
                    continue

                text = para.text.strip()
                style_name = para.style.name if para.style else ''

                # 检测列表
                if 'List' in style_name or 'Numbering' in style_name or \
                   child.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr') is not None:
                    markdown_content.append(f'- {text}')
                    continue

                # 检测标题
                if 'Heading' in style_name:
                    level = {'Heading 1': 1, 'Heading 2': 2, 'Heading 3': 3,
                             'Heading 4': 4, 'Heading 5': 5, 'Heading 6': 6}.get(style_name, 1)
                    if 'Heading 1' in style_name and (not title or title == file_title):
                        title = text
                    markdown_content.append(f"{'#' * level} {text}")
                    continue

                # 检测引用
                if 'Quote' in style_name or 'Block Text' in style_name:
                    markdown_content.append(f'> {text}')
                    continue

                # 检查段落中的图片
                drawings = child.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')

                if drawings:
                    # 有图片的段落
                    parts = []
                    for subchild in child:
                        subtag = subchild.tag.split('}')[-1] if '}' in subchild.tag else subchild.tag
                        if subtag == 'r':
                            drawing = subchild.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
                            if drawing is not None:
                                # 插入图片
                                if img_idx < len(saved_images) and saved_images[img_idx]:
                                    parts.append(f'\n![image]({saved_images[img_idx]})\n')
                                img_idx += 1
                            else:
                                t = subchild.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
                                if t is not None and t.text:
                                    parts.append(t.text)
                        elif subtag == 't' and subchild.text:
                            parts.append(subchild.text)

                    if parts:
                        markdown_content.append(''.join(parts))
                    elif not text:
                        # 纯图片段落
                        for _ in drawings:
                            if img_idx < len(saved_images) and saved_images[img_idx]:
                                markdown_content.append(f'\n![image]({saved_images[img_idx]})\n')
                            img_idx += 1
                else:
                    # 普通段落
                    if text:
                        parts = []
                        for run in para.runs:
                            txt = run.text or ''
                            try:
                                if run._element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}strike') is not None:
                                    txt = f'~~{txt}~~'
                            except Exception:
                                pass
                            if run.underline:
                                txt = f'<u>{txt}</u>'
                            if run.bold:
                                txt = f'**{txt}**'
                            if run.italic:
                                txt = f'*{txt}*'
                            parts.append(txt)
                        if parts:
                            markdown_content.append(''.join(parts))
                    else:
                        markdown_content.append('')

            elif tag == 'tbl':
                # 表格
                table = next((t for t in doc.tables if t._element is child), None)
                if table:
                    md_table = convert_table_to_markdown(table)
                    if md_table:
                        markdown_content.append(md_table)

        # 合并内容
        final_content = '\n\n'.join(markdown_content)

        return jsonify({
            'success': True,
            'title': title,
            'content': final_content
        })

    except Exception as e:
        return jsonify({'success': False, 'message': f'导入失败: {str(e)}'}), 500


def convert_table_to_markdown(table):
    """将Word表格转换为Markdown表格"""
    rows = []
    for i, row in enumerate(table.rows):
        cells = [cell.text.strip().replace('\n', ' ').replace('|', '\\|') for cell in row.cells]
        row_str = '| ' + ' | '.join(cells) + ' |'
        rows.append(row_str)

        # 第一行作为表头，加分隔符
        if i == 0:
            separator = '| ' + ' | '.join(['---'] * len(cells)) + ' |'
            rows.append(separator)

    return '\n'.join(rows) if rows else ''


@app.route('/api/admin/posts', methods=['POST'])
@login_required
def create_post():
    """创建文章"""
    data = request.get_json()
    
    conn = get_db()
    cursor = conn.cursor()
    
    # 生成 slug
    slug = data.get('slug') or data['title'].lower().replace(' ', '-')
    # 确保 slug 唯一
    cursor.execute('SELECT COUNT(*) FROM posts WHERE slug LIKE ?', (slug + '%',))
    if cursor.fetchone()[0] > 0:
        slug = f"{slug}-{cursor.fetchone()[0] + 1}"
    
    now = datetime.now().isoformat()
    
    cursor.execute('''
        INSERT INTO posts (title, slug, excerpt, content, tags, author, featured, status, created_at, updated_at)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
    ''', (data['title'], slug, data.get('excerpt', ''), data['content'], 
          data.get('tags', ''), data.get('author', '觉醒者'), 
          data.get('featured', 0), data.get('status', 'published'), now, now))
    
    post_id = cursor.lastrowid
    
    # 处理标签
    tags = data.get('tags', '').split(',')
    for tag_name in tags:
        tag_name = tag_name.strip()
        if tag_name:
            tag_slug = tag_name.lower()
            cursor.execute('INSERT OR IGNORE INTO tags (name, slug) VALUES (?, ?)', (tag_name, tag_slug))
            cursor.execute('SELECT id FROM tags WHERE slug = ?', (tag_slug,))
            tag_row = cursor.fetchone()
            if tag_row:
                cursor.execute('INSERT OR IGNORE INTO post_tags (post_id, tag_id) VALUES (?, ?)', 
                             (post_id, tag_row['id']))
    
    conn.commit()
    conn.close()
    
    return jsonify({'success': True, 'slug': slug})

@app.route('/api/admin/posts/<int:post_id>', methods=['PUT'])
@login_required
def update_post(post_id):
    """更新文章"""
    data = request.get_json()
    
    conn = get_db()
    cursor = conn.cursor()
    
    now = datetime.now().isoformat()
    
    cursor.execute('''
        UPDATE posts 
        SET title = ?, excerpt = ?, content = ?, tags = ?, featured = ?, status = ?, updated_at = ?
        WHERE id = ?
    ''', (data['title'], data.get('excerpt', ''), data['content'], 
          data.get('tags', ''), data.get('featured', 0), data.get('status', 'published'), 
          now, post_id))
    
    # 更新标签关联
    cursor.execute('DELETE FROM post_tags WHERE post_id = ?', (post_id,))
    tags = data.get('tags', '').split(',')
    for tag_name in tags:
        tag_name = tag_name.strip()
        if tag_name:
            tag_slug = tag_name.lower()
            cursor.execute('INSERT OR IGNORE INTO tags (name, slug) VALUES (?, ?)', (tag_name, tag_slug))
            cursor.execute('SELECT id FROM tags WHERE slug = ?', (tag_slug,))
            tag_row = cursor.fetchone()
            if tag_row:
                cursor.execute('INSERT OR IGNORE INTO post_tags (post_id, tag_id) VALUES (?, ?)', 
                             (post_id, tag_row['id']))
    
    conn.commit()
    conn.close()
    
    return jsonify({'success': True})

@app.route('/api/admin/posts/<int:post_id>', methods=['DELETE'])
@login_required
def delete_post(post_id):
    """删除文章"""
    conn = get_db()
    cursor = conn.cursor()
    cursor.execute('DELETE FROM post_tags WHERE post_id = ?', (post_id,))
    cursor.execute('DELETE FROM posts WHERE id = ?', (post_id,))
    conn.commit()
    conn.close()
    return jsonify({'success': True})

@app.route('/api/admin/posts', methods=['GET'])
@login_required
def get_all_posts():
    """获取所有文章（管理后台）"""
    conn = get_db()
    cursor = conn.cursor()
    cursor.execute('SELECT * FROM posts ORDER BY created_at DESC')
    posts = [dict(row) for row in cursor.fetchall()]
    conn.close()
    return jsonify(posts)

# ============================================
# 静态文件
# ============================================

@app.route('/static/<path:filename>')
def serve_static(filename):
    """提供静态文件服务"""
    from flask import send_from_directory
    import os
    static_dir = os.path.join(os.path.dirname(__file__), 'static')
    return send_from_directory(static_dir, filename)

# ============================================
# 页面路由
# ============================================

@app.route('/')
def index():
    """首页"""
    settings = get_settings()
    return render_template('index.html', settings=settings, version=VERSION)

@app.route('/admin')
def admin():
    """管理后台"""
    if 'admin' not in session:
        return redirect(url_for('admin_login'))
    settings = get_settings()
    return render_template('admin.html', settings=settings, version=VERSION)

@app.route('/admin/login')
def admin_login():
    """登录页"""
    if 'admin' in session:
        return redirect(url_for('admin'))
    settings = get_settings()
    return render_template('login.html', settings=settings, version=VERSION)

@app.route('/tags')
def tags_page():
    """标签页面（重定向到 SPA）"""
    return redirect('/#tags')

@app.route('/about')
def about():
    """关于页面"""
    settings = get_settings()
    conn = get_db()
    cursor = conn.cursor()
    cursor.execute('SELECT content FROM about_page WHERE id = 1')
    row = cursor.fetchone()
    conn.close()
    
    content = row['content'] if row else ''
    # 将 Markdown 转换为 HTML
    about_html = markdown.markdown(content, extensions=['fenced_code', 'tables'])
    
    return render_template('about.html', settings=settings, about_content=about_html, version=VERSION)

# ============================================
# 静态文件
# ============================================

@app.route('/css/<path:filename>')
def serve_css(filename):
    response = send_from_directory('templates/css', filename)
    response.headers['Cache-Control'] = 'public, max-age=31536000'
    return response

@app.route('/js/<path:filename>')
def serve_js(filename):
    response = send_from_directory('templates/js', filename)
    response.headers['Cache-Control'] = 'public, max-age=31536000'
    return response

# ============================================
# 初始化
# ============================================

if __name__ == '__main__':
    init_db()
    print("=" * 50)
    print("认知觉醒博客系统已启动")
    print("前台地址: http://0.0.0.0:5000")
    print("后台地址: http://0.0.0.0:5000/admin")
    print("默认账号: admin / admin123")
    print("=" * 50)
    app.run(debug=True, host='0.0.0.0', port=5000)
